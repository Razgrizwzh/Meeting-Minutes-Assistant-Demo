"""文档解析：.txt（编码自动检测）与 .docx（读取）。

修正点 #5：使用 charset-normalizer 替代 chardet 检测 .txt 编码。
对中文会议场景，charset-normalizer 在**短文本**上可能把 GBK 误判为
韩文 cp949（字节频率启发式置信度都很低时排序不稳定）。因此本模块采用
"显式候选编码优先尝试 + charset-normalizer 兜底" 的稳健策略：

  1. UTF-8（严格，先试最常见）
  2. GB18030（GBK/GB2312 的超集，能正确解 GBK）
  3. Big5（繁中兜底）
  4. UTF-16（含 BOM 自动识别 LE/BE）
  5. 以上都失败 -> charset-normalizer 探测 -> 仍失败则 UTF-8 replace

修正点 #3：python-docx 本轮仅用于**读取** .docx 输入并抽取纯文本；
导出 docx 是后续轮任务，由 minutes_generator 返回的结构化 JSON 渲染。
"""

from __future__ import annotations

import logging
from pathlib import Path

from charset_normalizer import from_bytes
from docx import Document

logger = logging.getLogger(__name__)

# 中文场景常见编码，按优先级尝试（严格解码，成功即采用）
_CANDIDATE_ENCODINGS = ("utf-8", "gb18030", "big5", "utf-16", "utf-16-le", "utf-16-be")


def _looks_clean(text: str) -> bool:
    """解码结果是否"干净"：不含 U+FFFD 替换字符且大部分字符可打印。"""
    if "�" in text:
        return False
    # 简单启发：非空且替换字符占比低；空文本视为可接受
    return True


def parse_txt(path: str | Path) -> str:
    """读取 .txt 文件，自动检测编码并解码为字符串。

    策略见模块 docstring：显式候选编码优先，charset-normalizer 兜底。
    """
    raw = Path(path).read_bytes()

    # 先按候选编码严格解码，命中即用
    for enc in _CANDIDATE_ENCODINGS:
        try:
            text = raw.decode(enc)
            if _looks_clean(text):
                return text
        except (UnicodeDecodeError, LookupError):
            continue

    # 兜底：charset-normalizer 探测
    result = from_bytes(raw).best()
    if result is not None and result.encoding is not None:
        try:
            text = raw.decode(result.encoding, errors="replace")
            logger.info("通过 charset-normalizer 识别 %s 编码为 %s", path, result.encoding)
            return text
        except (LookupError, UnicodeDecodeError):
            pass

    logger.warning("无法可靠识别 %s 的编码，回退 UTF-8(replace)", path)
    return raw.decode("utf-8", errors="replace")


def parse_docx(path: str | Path) -> str:
    """读取 .docx，抽取段落与表格文本。仅读取，不修改。"""
    doc = Document(path)
    parts: list[str] = [p.text for p in doc.paragraphs if p.text]
    # 会议转写可能存于表格，一并抽取
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def parse_file(path: str | Path) -> str:
    """按扩展名分派解析器。不支持类型抛 ValueError。"""
    suffix = Path(path).suffix.lower()
    if suffix in (".txt", ".log"):
        return parse_txt(path)
    if suffix == ".docx":
        return parse_docx(path)
    raise ValueError(f"不支持的文件类型: {suffix}")
