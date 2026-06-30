"""docx 导出：由结构化纪要 JSON 渲染 Word 文档。

【修正点 #3 —— 不解析 Markdown】
数据源是结构化 JSON，用 python-docx 直接渲染（add_heading/add_paragraph/add_table），
无需任何 Markdown 解析。表格用内置 'Table Grid' 样式（开箱可用）。
"""

from __future__ import annotations

from pathlib import Path

from docx import Document


def render_docx(minutes: dict, output_path: str | Path) -> Path:
    """将结构化纪要 JSON 渲染为 .docx，写入 output_path 并返回该路径。"""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    name = minutes.get("meeting_name", "未命名会议")
    date = minutes.get("date", "")
    attendees = minutes.get("attendees", []) or []
    topics = minutes.get("topics", []) or []
    decisions = minutes.get("decisions", []) or []
    todos = minutes.get("todos", []) or []

    doc = Document()
    doc.add_heading("会议纪要", level=1)

    p = doc.add_paragraph()
    p.add_run("会议名称：").bold = True
    p.add_run(name)
    p = doc.add_paragraph()
    p.add_run("日期：").bold = True
    p.add_run(date or "—")
    p = doc.add_paragraph()
    p.add_run("参会人员：").bold = True
    p.add_run("、".join(attendees) if attendees else "—")

    doc.add_heading("核心议题与讨论要点", level=2)
    if topics:
        for i, t in enumerate(topics, 1):
            doc.add_heading(f"{i}. {t.get('title', '（无标题）')}", level=3)
            for pt in t.get("points", []) or []:
                doc.add_paragraph(pt, style="List Bullet")
    else:
        doc.add_paragraph("（无）")

    doc.add_heading("决策结论", level=2)
    if decisions:
        for d in decisions:
            text = d.get("text", "")
            owner = d.get("owner", "")
            line = f"{text}（负责人：{owner}）" if owner else text
            doc.add_paragraph(line, style="List Bullet")
    else:
        doc.add_paragraph("（无）")

    doc.add_heading("待办事项", level=2)
    if todos:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "事项", "负责人", "截止时间"
        for cell in hdr:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
        for td in todos:
            row = table.add_row().cells
            row[0].text = td.get("item", "")
            row[1].text = td.get("owner", "") or "—"
            row[2].text = td.get("due", "") or "—"
    else:
        doc.add_paragraph("（无）")

    doc.save(str(output_path))
    return output_path